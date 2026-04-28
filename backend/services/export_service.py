import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import logging
from datetime import datetime

class ExportService:
    def __init__(self):
        print("DEBUG: ExportService class initialized")
        # Base exports directory - use project relative path to avoid permission issues
        # This will create an 'exports' folder inside your AI_Exam_Generator directory
        current_dir = os.path.dirname(os.path.abspath(__file__)) # backend/services
        backend_dir = os.path.dirname(current_dir) # backend
        project_root = os.path.dirname(backend_dir) # AI_Exam_Generator
        
        self.output_dir = os.path.join(project_root, "exports")
        try:
            os.makedirs(self.output_dir, exist_ok=True)
        except Exception as e:
            logging.error(f"Failed to create base export directory: {e}")
            # Fallback to a temporary directory if project root is read-only
            import tempfile
            self.output_dir = os.path.join(tempfile.gettempdir(), "ai_exam_exports")
            os.makedirs(self.output_dir, exist_ok=True)

    def _sanitize_filename(self, name: str) -> str:
        import re
        if not name or str(name).lower() == "assessment" or str(name).lower() == "exam":
            return "" # Return empty if it's just a placeholder
        # Replace spaces and special chars with underscores
        name = re.sub(r'[^a-zA-Z0-9]', '_', str(name).strip())
        # Remove multiple underscores
        name = re.sub(r'_+', '_', name).strip('_')
        return name

    def _get_export_path(self, subject: str, format_type: str) -> str:
        """
        Creates and returns path: /exports/Subject/PDF or /exports/Subject/DOCX
        """
        clean_subject = self._sanitize_filename(subject) or "General"
        # Capitalize first letter for folder name
        folder_name = clean_subject.capitalize()
        path = os.path.join(self.output_dir, folder_name, format_type.upper())
        os.makedirs(path, exist_ok=True)
        return path

    def _generate_clean_filename(self, subject: str, topic: str, doc_type: str, ext: str) -> str:
        """
        Generates filename: Subject_Topic_Type.ext
        """
        s = self._sanitize_filename(subject)
        t = self._sanitize_filename(topic)
        
        parts = []
        if s: parts.append(s)
        if t: parts.append(t)
        parts.append(doc_type) # e.g. Exam, Paper, Quiz
        
        return "_".join(parts) + f".{ext}"

    def generate_docx_file(self, session_id: str, questions: list, branding=None, student_info=None, is_answer_key=False, include_answers=True, subject="Exam", topic="Assessment", exam_title="Final Examination", time_limit="2 Hours", total_marks=100, passing_marks=40, passing_percentage=40, mcq_marks=1, short_marks=4, long_marks=10, prog_marks=15) -> str:
        doc = Document()
        
        # Header Table for Logo and Title
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        
        # Logo (Left)
        if branding and branding.logo_path and os.path.exists(branding.logo_path):
            try:
                cell = table.cell(0, 0)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(branding.logo_path, width=Pt(60))
            except Exception as e:
                logging.error(f"Error adding logo to DOCX: {e}")

        # Title (Center)
        cell = table.cell(0, 1)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(exam_title.upper())
        run.bold = True
        run.font.size = Pt(20)
        
        # Branding Info (Right)
        if branding and branding.uni:
            cell = table.cell(0, 2)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(branding.uni)
            run.bold = True
            run.font.size = Pt(10)
            if branding.dept:
                paragraph.add_run(f"\n{branding.dept}")

        doc.add_paragraph() # Spacer
        
        # Student Info Section (if enabled)
        if student_info and getattr(student_info, 'enabled', False) and not is_answer_key:
            # We use a table for better alignment of student info
            s_table = doc.add_table(rows=0, cols=2)
            s_table.autofit = True
            
            # Row 1: Name and Roll No
            if getattr(student_info, 'show_name', True) or getattr(student_info, 'show_roll_no', True):
                row = s_table.add_row()
                if getattr(student_info, 'show_name', True):
                    p = row.cells[0].paragraphs[0]
                    p.add_run("STUDENT NAME: ").bold = True
                    p.add_run("." * 40)
                if getattr(student_info, 'show_roll_no', True):
                    p = row.cells[1].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p.add_run("ROLL NO: ").bold = True
                    p.add_run("." * 20)

            # Row 2: Class and Section/Date
            if getattr(student_info, 'show_class', True) or getattr(student_info, 'show_date', True) or getattr(student_info, 'show_section', True):
                row = s_table.add_row()
                if getattr(student_info, 'show_class', True):
                    p = row.cells[0].paragraphs[0]
                    p.add_run("CLASS / SEMESTER: ").bold = True
                    p.add_run("." * 30)
                
                right_cell = row.cells[1]
                p = right_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                if getattr(student_info, 'show_section', True):
                    p.add_run("SECTION: ").bold = True
                    p.add_run("." * 10)
                    if getattr(student_info, 'show_date', True):
                        p.add_run("  DATE: ").bold = True
                        p.add_run("." * 15)
                elif getattr(student_info, 'show_date', True):
                    p.add_run("DATE: ").bold = True
                    p.add_run("." * 15)
            
            doc.add_paragraph() # Spacer after student info

        # Metadata Row
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.add_run(f"Subject: {subject} | Time: {time_limit}\n").bold = True
        meta_p.add_run(f"Total Marks: {total_marks} | Passing Marks: {passing_marks} ({passing_percentage}%)").bold = True
        
        doc.add_paragraph("-" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

        if is_answer_key:
            doc.add_paragraph("ANSWER KEY").alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._write_docx_answer_key(doc, questions, student_info)
        else:
            self._write_docx_question_paper(doc, questions, include_answers, mcq_marks, short_marks, long_marks, prog_marks, student_info)

        # File Naming & Folder Structure
        export_dir = self._get_export_path(subject, "DOCX")
        filename = f"{subject}_{total_marks}Marks.{'docx' if not is_answer_key else 'key.docx'}"
        filename = self._sanitize_filename(filename) + ".docx"
        
        file_path = os.path.join(export_dir, filename)
        doc.save(file_path)
        return file_path

    def _write_docx_question_paper(self, doc, questions, include_answers, mcq_marks, short_marks, long_marks, prog_marks, student_info=None):
        mcqs = [q for q in questions if q['type'].lower() == 'mcq']
        shorts = [q for q in questions if q['type'].lower() == 'short']
        longs = [q for q in questions if q['type'].lower() == 'long']
        progs = [q for q in questions if q['type'].lower() == 'programming']

        q_count = 1
        show_bloom = student_info and getattr(student_info, 'show_bloom_tags', False)

        if mcqs:
            doc.add_paragraph(f"SECTION A: Multiple Choice Questions ({len(mcqs)} × {mcq_marks} = {len(mcqs)*mcq_marks} marks)").bold = True
            doc.add_paragraph("Instruction: Attempt all questions. Choose the correct option.").italic = True
            
            multi_col = student_info and getattr(student_info, 'multi_column_mcqs', False)
            
            if multi_col:
                mcq_table = doc.add_table(rows=0, cols=2)
                mcq_table.autofit = True
                for i in range(0, len(mcqs), 2):
                    row = mcq_table.add_row()
                    q1 = mcqs[i]
                    p1 = row.cells[0].paragraphs[0]
                    bloom_tag = f"[{q1.get('bloom_level', 'Apply')}] " if show_bloom else ""
                    p1.add_run(f"Q{q_count}. ").bold = True
                    p1.add_run(f"{bloom_tag}{q1['question']} ({mcq_marks} mark)")
                    if 'options' in q1:
                        for opt in q1['options']:
                            row.cells[0].add_paragraph(f"   {opt}")
                    q_count += 1
                    
                    if i + 1 < len(mcqs):
                        q2 = mcqs[i+1]
                        p2 = row.cells[1].paragraphs[0]
                        bloom_tag = f"[{q2.get('bloom_level', 'Apply')}] " if show_bloom else ""
                        p2.add_run(f"Q{q_count}. ").bold = True
                        p2.add_run(f"{bloom_tag}{q2['question']} ({mcq_marks} mark)")
                        if 'options' in q2:
                            for opt in q2['options']:
                                row.cells[1].add_paragraph(f"   {opt}")
                        q_count += 1
            else:
                for q in mcqs:
                    p = doc.add_paragraph()
                    bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
                    p.add_run(f"Q{q_count}. ").bold = True
                    p.add_run(f"{bloom_tag}{q['question']} ({mcq_marks} mark)")
                    if 'options' in q:
                        for opt in q['options']:
                            doc.add_paragraph(f"   {opt}")
                    q_count += 1
            doc.add_paragraph()

        if shorts:
            doc.add_paragraph(f"SECTION B: Short Answer Questions ({len(shorts)} × {short_marks} = {len(shorts)*short_marks} marks)").bold = True
            doc.add_paragraph("Instruction: Attempt all questions. Answer briefly.").italic = True
            for q in shorts:
                p = doc.add_paragraph()
                bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
                p.add_run(f"Q{q_count}. ").bold = True
                p.add_run(f"{bloom_tag}{q['question']} ({short_marks} marks)")
                q_count += 1
            doc.add_paragraph()

        if longs:
            doc.add_paragraph(f"SECTION C: Long Answer Questions ({len(longs)} × {long_marks} = {len(longs)*long_marks} marks)").bold = True
            doc.add_paragraph("Instruction: Attempt all questions. Answer in detail.").italic = True
            for q in longs:
                p = doc.add_paragraph()
                bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
                p.add_run(f"Q{q_count}. ").bold = True
                p.add_run(f"{bloom_tag}{q['question']} ({long_marks} marks)")
                q_count += 1
            doc.add_paragraph()

        if progs:
            doc.add_paragraph(f"SECTION D: Programming Questions ({len(progs)} × {prog_marks} = {len(progs)*prog_marks} marks)").bold = True
            doc.add_paragraph("Instruction: Attempt all questions. Write clear and efficient code.").italic = True
            for q in progs:
                p = doc.add_paragraph()
                bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
                p.add_run(f"Q{q_count}. ").bold = True
                p.add_run(f"{bloom_tag}{q['question']} ({prog_marks} marks)")
                q_count += 1
            doc.add_paragraph()

    def _write_docx_answer_key(self, doc, questions, student_info=None):
        q_count = 1
        show_bloom = student_info and getattr(student_info, 'show_bloom_tags', False)
        for q in questions:
            p = doc.add_paragraph()
            bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
            p.add_run(f"Q{q_count}: ").bold = True
            p.add_run(f"{bloom_tag}{q['question']}")
            
            ans_p = doc.add_paragraph()
            ans_p.add_run("CORRECT ANSWER: ").bold = True
            ans_p.add_run(str(q['answer']))
            doc.add_paragraph()
            q_count += 1

    async def generate_pdf_file(self, session_id: str, questions: list, branding=None, student_info=None, is_answer_key=False, include_answers=True, subject="Exam", topic="Assessment", exam_title="Final Examination", time_limit="2 Hours", total_marks=100, passing_marks=40, passing_percentage=40, mcq_marks=1, short_marks=4, long_marks=10, prog_marks=15) -> str:
        try:
            watermark_text = branding.watermark_text if branding and getattr(branding, 'enable_watermark', False) else None
            pdf = PDFWithWatermark(watermark_text=watermark_text)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            
            # Logo & Header
            y_start = pdf.get_y()
            if branding and branding.logo_path and os.path.exists(branding.logo_path):
                try:
                    pdf.image(branding.logo_path, 10, y_start, 25)
                except Exception as e:
                    logging.error(f"Error adding logo to PDF: {e}")
            
            effective_width = pdf.w - pdf.l_margin - pdf.r_margin - 2 # 2mm safety margin
            pdf.set_font("Arial", 'B', 20)
            pdf.cell(effective_width, 10, exam_title.upper(), ln=True, align='C')
            
            pdf.set_font("Arial", 'B', 12)
            if branding and branding.uni:
                pdf.cell(effective_width, 8, branding.uni, ln=True, align='C')
                if branding.dept:
                    pdf.set_font("Arial", '', 10)
                    pdf.cell(effective_width, 6, branding.dept, ln=True, align='C')
            
            pdf.ln(5)
            
            # Student Info Section (if enabled)
            if student_info and getattr(student_info, 'enabled', False) and not is_answer_key:
                pdf.set_font("Arial", 'B', 10)
                pdf.ln(2)
                
                has_name = getattr(student_info, 'show_name', True)
                has_roll = getattr(student_info, 'show_roll_no', True)
                
                if has_name:
                    pdf.cell(110, 8, "STUDENT NAME: " + "." * 40, ln=0)
                if has_roll:
                    if not has_name: pdf.set_x(10)
                    else: pdf.set_x(130)
                    pdf.cell(70, 8, "ROLL NO: " + "." * 20, ln=1, align='R')
                elif has_name:
                    pdf.ln(8)

                has_class = getattr(student_info, 'show_class', True)
                has_date = getattr(student_info, 'show_date', True)
                has_section = getattr(student_info, 'show_section', True)

                if has_class or has_date or has_section:
                    if has_class:
                        pdf.cell(80, 8, "CLASS / SEMESTER: " + "." * 25, ln=0)
                    
                    if has_section or has_date:
                        pdf.set_x(pdf.w - pdf.r_margin - 100) # Ensure it stays within right margin
                        section_text = ""
                        if has_section:
                            section_text += "SECTION: " + "." * 10
                        if has_date:
                            if section_text: section_text += "   "
                            section_text += "DATE: " + "." * 15
                        pdf.cell(100, 8, section_text, ln=1, align='R')
                    else:
                        pdf.ln(8)
                
                pdf.ln(4)

            pdf.set_font("Arial", 'B', 10)
            pdf.cell(effective_width, 8, f"Subject: {subject} | Time: {time_limit}", ln=True, align='C')
            pdf.cell(effective_width, 8, f"Total Marks: {total_marks} | Passing Marks: {passing_marks} ({passing_percentage}%)", ln=True, align='C')
            
            pdf.ln(2)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(8)
            
            if is_answer_key:
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(effective_width, 10, "ANSWER KEY", ln=True, align='C')
                pdf.ln(5)
                self._write_pdf_answer_key(pdf, questions, effective_width, student_info)
            else:
                self._write_pdf_question_paper(pdf, questions, effective_width, include_answers, mcq_marks, short_marks, long_marks, prog_marks, student_info)

            export_dir = self._get_export_path(subject, "PDF")
            filename = f"{subject}_{total_marks}Marks.{'pdf' if not is_answer_key else 'key.pdf'}"
            filename = self._sanitize_filename(filename) + ".pdf"
            
            file_path = os.path.join(export_dir, filename)
            pdf.output(file_path)
            return file_path
        except Exception as e:
            logging.error(f"PDF error: {e}")
            raise e

    def _clean_text_for_pdf(self, text: str) -> str:
        if not text: return ""
        replacements = {'→': '->', '←': '<-', '•': '*', '—': '-', '–': '-', '“': '"', '”': '"', '‘': "'", '’': "'", '≥': '>=', '≤': '<=', '≠': '!=', '±': '+/-'}
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        return text.encode('latin-1', 'replace').decode('latin-1')

    def _write_pdf_question_paper(self, pdf, questions, effective_width, include_answers, mcq_marks, short_marks, long_marks, prog_marks, student_info=None):
        mcqs = [q for q in questions if q['type'].lower() == 'mcq']
        shorts = [q for q in questions if q['type'].lower() == 'short']
        longs = [q for q in questions if q['type'].lower() == 'long']
        progs = [q for q in questions if q['type'].lower() == 'programming']

        q_count = 1
        show_bloom = student_info and getattr(student_info, 'show_bloom_tags', False)
        multi_col = student_info and getattr(student_info, 'multi_column_mcqs', False)

        if mcqs:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, f"SECTION A: Multiple Choice Questions ({len(mcqs)} x {mcq_marks} = {len(mcqs)*mcq_marks} marks)", ln=True)
            pdf.set_font("Arial", 'I', 10)
            pdf.cell(0, 8, "Instruction: Attempt all questions. Choose the correct option.", ln=True)
            pdf.ln(2)
            
            if multi_col:
                col_width = effective_width / 2
                for i, q in enumerate(mcqs):
                    bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
                    pdf.set_font("Arial", 'B', 10)
                    q_text = self._clean_text_for_pdf(f"Q{q_count}. {bloom_tag}{q['question']} ({mcq_marks})")
                    current_y = pdf.get_y()
                    current_x = 10 if i % 2 == 0 else 10 + col_width
                    pdf.set_xy(current_x, current_y)
                    pdf.multi_cell(col_width - 5, 6, q_text)
                    pdf.set_font("Arial", '', 9)
                    if 'options' in q:
                        for opt in q['options']:
                            pdf.set_x(current_x + 5)
                            pdf.multi_cell(col_width - 10, 5, self._clean_text_for_pdf(opt))
                    if i % 2 != 0 or i == len(mcqs) - 1:
                        pdf.ln(4)
                    else:
                        pdf.set_y(current_y)
                    q_count += 1
            else:
                for q in mcqs:
                    pdf.set_font("Arial", 'B', 11)
                    bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
                    pdf.multi_cell(effective_width, 7, self._clean_text_for_pdf(f"Q{q_count}. {bloom_tag}{q['question']} ({mcq_marks} mark)"))
                    pdf.set_font("Arial", '', 10)
                    if 'options' in q:
                        for opt in q['options']:
                            pdf.set_x(20)
                            # Adjust width for indentation (20 - 10 = 10mm indent)
                            pdf.multi_cell(effective_width - 10, 6, self._clean_text_for_pdf(opt))
                    q_count += 1
                    pdf.ln(4)
            pdf.ln(5)

        for section_title, qs, marks, instr in [
            ("SECTION B: Short Answer Questions", shorts, short_marks, "Instruction: Attempt all questions. Answer briefly."),
            ("SECTION C: Long Answer Questions", longs, long_marks, "Instruction: Attempt all questions. Answer in detail."),
            ("SECTION D: Programming Questions", progs, prog_marks, "Instruction: Attempt all questions. Write clear and efficient code.")
        ]:
            if qs:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, f"{section_title} ({len(qs)} x {marks} = {len(qs)*marks} marks)", ln=True)
                pdf.set_font("Arial", 'I', 10)
                pdf.cell(0, 8, instr, ln=True)
                pdf.ln(2)
                for q in qs:
                    pdf.set_font("Arial", 'B', 11)
                    bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
                    pdf.multi_cell(effective_width, 7, self._clean_text_for_pdf(f"Q{q_count}. {bloom_tag}{q['question']} ({marks} marks)"))
                    q_count += 1
                    pdf.ln(4)
                pdf.ln(5)

    def _write_pdf_answer_key(self, pdf, questions, effective_width, student_info=None):
        q_count = 1
        show_bloom = student_info and getattr(student_info, 'show_bloom_tags', False)
        # Use a slightly narrower width for safety and indentation
        key_width = effective_width - 10
        for q in questions:
            pdf.set_font("Arial", 'B', 11)
            bloom_tag = f"[{q.get('bloom_level', 'Apply')}] " if show_bloom else ""
            q_text = self._clean_text_for_pdf(f"Q{q_count}: {bloom_tag}{q['question']}")
            pdf.multi_cell(effective_width, 7, q_text)
            
            pdf.set_font("Arial", 'I', 10)
            pdf.set_text_color(0, 128, 0)
            pdf.set_x(20) # Indent answer
            ans_text = f"CORRECT ANSWER: {self._clean_text_for_pdf(str(q['answer']))}"
            pdf.multi_cell(key_width, 7, ans_text)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(5)
            q_count += 1

    def cleanup_exports(self, session_id: str):
        pass

class PDFWithWatermark(FPDF):
    def __init__(self, watermark_text=None, logo_path=None):
        super().__init__()
        self.watermark_text = watermark_text
        self.watermark_logo = logo_path

    def header(self):
        if self.watermark_text:
            self.set_font('helvetica', 'B', 50)
            self.set_text_color(220, 220, 220) # Light grey
            with self.rotation(45, 105, 148):
                self.text(40, 190, self.watermark_text)
            self.set_text_color(0, 0, 0)
        
        if self.watermark_logo and os.path.exists(self.watermark_logo):
            self.image(self.watermark_logo, x=55, y=100, w=100)
